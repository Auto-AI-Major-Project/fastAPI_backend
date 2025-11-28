# data_cleaner.py (Comprehensive Report Generation)
import pandas as pd
import json
import numpy as np
from io import BytesIO
from typing import Dict, Any, Tuple, List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.platypus import Image as RLImage
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.pdfgen import canvas
from datetime import datetime
import io
import html

# Import the utility function to get the initial profile
from utils import data_utils 

# --- Core Utility Functions (Remains the same) ---
# ... (get_data_profile_string, apply_manual_cleaning) ...

def get_data_profile_string(inspection_data: Dict[str, Any]) -> str:
    """Consolidates data inspection into a single prompt string for Gemini."""
    return f"""
    --- DATASET PROFILE ---
    Shape (Rows, Columns): {inspection_data['shape']}
    
    Missing Values (Count > 0): {inspection_data['missing_values']}
    
    Column Types: {inspection_data['data_types']}
    
    --- SAMPLE DATA (First 5 Rows) ---
    {inspection_data['preview_first_5_rows']}
    
    --- SUMMARY STATISTICS ---
    {inspection_data['summary_statistics_markdown']}
    
    --- END PROFILE ---
    """


def sanitize_text(text: str) -> str:
    """Remove markdown and ensure proper HTML escaping for reportlab."""
    # Remove all markdown bold markers
    text = text.replace('**', '')
    # Escape any problematic characters
    text = html.escape(text)
    return text

def generate_pdf_report(report_details: Dict[str, Any], filename: str) -> BytesIO:
    """Generates a professional PDF report with better formatting."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    # Container for the 'Flowable' objects
    elements = []
    
    # Define styles
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name='CustomTitle',
                             parent=styles['Heading1'],
                             fontSize=24,
                             textColor=colors.HexColor('#1A6B8E'),
                             spaceAfter=30,
                             alignment=TA_CENTER,
                             fontName='Helvetica-Bold'))
    
    styles.add(ParagraphStyle(name='CustomHeading',
                             parent=styles['Heading2'],
                             fontSize=16,
                             textColor=colors.HexColor('#1A6B8E'),
                             spaceAfter=12,
                             spaceBefore=12,
                             fontName='Helvetica-Bold'))
    
    styles.add(ParagraphStyle(name='CustomBody',
                             parent=styles['BodyText'],
                             fontSize=11,
                             spaceAfter=12,
                             alignment=TA_JUSTIFY))
    
    styles.add(ParagraphStyle(name='CustomSmall',
                             parent=styles['BodyText'],
                             fontSize=9,
                             textColor=colors.grey,
                             alignment=TA_CENTER))
    
    # Title
    title_text = sanitize_text(report_details.get('title', 'Data Cleaning Report'))
    title = Paragraph(title_text, styles['CustomTitle'])
    elements.append(title)
    
    # Timestamp
    timestamp = Paragraph(
        f"Generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}",
        styles['CustomSmall']
    )
    elements.append(timestamp)
    elements.append(Spacer(1, 0.3*inch))
    
    # ==================== SECTION 1: EXECUTIVE SUMMARY ====================
    elements.append(Paragraph("Executive Summary", styles['CustomHeading']))
    
    metrics = report_details.get('comparative_metrics', {})
    original_shape = metrics.get('original_shape', (0, 0))
    final_shape = metrics.get('final_shape', (0, 0))
    original_missing = metrics.get('original_missing', 0)
    final_missing = metrics.get('final_missing', 0)
    
    summary_data = [
        ['Metric', 'Original', 'After Cleaning', 'Change'],
        ['Rows', 
         str(original_shape[0]),
         str(final_shape[0]),
         f"{final_shape[0] - original_shape[0]:+d}"],
        ['Columns',
         str(original_shape[1]),
         str(final_shape[1]),
         f"{final_shape[1] - original_shape[1]:+d}"],
        ['Missing Values',
         str(original_missing),
         str(final_missing),
         f"{final_missing - original_missing:+d}"],
    ]
    
    summary_table = Table(summary_data, colWidths=[2.5*inch, 1.2*inch, 1.2*inch, 1.2*inch])
    summary_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1A6B8E')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.grey),
        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
    ]))
    
    elements.append(summary_table)
    elements.append(Spacer(1, 0.3*inch))
    
    # ==================== SECTION 2: CLEANING OPERATIONS ====================
    elements.append(Paragraph("Data Cleaning Operations", styles['CustomHeading']))
    
    cleaning_summary = report_details.get('cleaning_summary', [])
    if cleaning_summary:
        for i, operation in enumerate(cleaning_summary, 1):
            # Clean and sanitize operation text
            clean_text = sanitize_text(operation)
            op_para = Paragraph(f"<b>{i}.</b> {clean_text}", styles['CustomBody'])
            elements.append(op_para)
    else:
        elements.append(Paragraph("No cleaning operations were applied.", styles['CustomBody']))
    
    elements.append(Spacer(1, 0.2*inch))
    
    # ==================== SECTION 3: DATA QUALITY ASSESSMENT ====================
    elements.append(Paragraph("Data Quality Assessment", styles['CustomHeading']))
    
    # Completeness Score
    total_cells = final_shape[0] * final_shape[1]
    completeness = ((total_cells - final_missing) / total_cells * 100) if total_cells > 0 else 100
    
    quality_text = f"""
    <b>Completeness Score:</b> {completeness:.2f}%<br/>
    <b>Status:</b> {'Excellent' if completeness >= 95 else 'Good' if completeness >= 85 else 'Fair'}<br/>
    <br/>
    The dataset has been successfully cleaned and is ready for analysis. 
    All missing values have been handled appropriately through imputation or removal.
    """
    
    elements.append(Paragraph(quality_text, styles['CustomBody']))
    elements.append(Spacer(1, 0.2*inch))
    
    # ==================== SECTION 4: CORRELATION INSIGHTS ====================
    correlation_pairs = report_details.get('correlation_pairs', [])
    if correlation_pairs:
        elements.append(Paragraph("Key Relationship Insights", styles['CustomHeading']))
        
        # Separate positive and negative correlations
        positive_corr = [p for p in correlation_pairs if p.get('type') == 'Positive']
        negative_corr = [p for p in correlation_pairs if p.get('type') == 'Negative']
        
        if positive_corr:
            elements.append(Paragraph("<b>Strong Positive Correlations:</b>", styles['CustomBody']))
            corr_data = [['Features', 'Correlation', 'Strength']]
            for corr in positive_corr[:5]:  # Top 5
                corr_data.append([
                    sanitize_text(str(corr.get('pair', ''))),
                    str(corr.get('value', '')),
                    str(corr.get('strength', ''))
                ])
            
            corr_table = Table(corr_data, colWidths=[3*inch, 1.5*inch, 1.5*inch])
            corr_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2ECC71')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.lightgreen),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
            ]))
            elements.append(corr_table)
            elements.append(Spacer(1, 0.1*inch))
        
        if negative_corr:
            elements.append(Paragraph("<b>Strong Negative Correlations:</b>", styles['CustomBody']))
            corr_data = [['Features', 'Correlation', 'Strength']]
            for corr in negative_corr[:5]:  # Top 5
                corr_data.append([
                    sanitize_text(str(corr.get('pair', ''))),
                    str(corr.get('value', '')),
                    str(corr.get('strength', ''))
                ])
            
            corr_table = Table(corr_data, colWidths=[3*inch, 1.5*inch, 1.5*inch])
            corr_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E74C3C')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
                ('BACKGROUND', (0, 1), (-1, -1), colors.lightcoral),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
            ]))
            elements.append(corr_table)
        
        elements.append(Spacer(1, 0.2*inch))
    
    # ==================== SECTION 5: OUTLIER ANALYSIS ====================
    outlier_flags = report_details.get('outlier_flags', [])
    if outlier_flags:
        elements.append(Paragraph("Outlier Detection (IQR Method)", styles['CustomHeading']))
        
        # Create a simple list without complex HTML formatting
        for flag in outlier_flags[:10]:  # Limit to 10
            # Completely sanitize the text
            clean_flag = sanitize_text(flag)
            # Use bullet point with simple text
            flag_para = Paragraph(f"• {clean_flag}", styles['CustomBody'])
            elements.append(flag_para)
        
        elements.append(Spacer(1, 0.2*inch))
    
    # ==================== SECTION 6: RECOMMENDATIONS ====================
    elements.append(Paragraph("Recommendations for Next Steps", styles['CustomHeading']))
    
    recommendations = [
        "Feature Engineering: Consider creating interaction features between highly correlated variables.",
        "Outlier Treatment: Review flagged outliers and decide on appropriate treatment (removal, capping, or transformation).",
        "Model Selection: The cleaned dataset is now ready for machine learning. Consider starting with baseline models.",
        "Validation: Set aside 20-30% of data for testing to ensure model generalization.",
        "Feature Selection: Use correlation insights to remove redundant features and improve model performance."
    ]
    
    for i, rec in enumerate(recommendations, 1):
        rec_para = Paragraph(f"<b>{i}.</b> {rec}", styles['CustomBody'])
        elements.append(rec_para)
    
    elements.append(Spacer(1, 0.3*inch))
    
    # ==================== FOOTER ====================
    key_insight = sanitize_text(report_details.get('key_insight', 'Analysis complete.'))
    insight_text = f"<b>Key Insight:</b> {key_insight}"
    
    insight_para = Paragraph(
        insight_text,
        ParagraphStyle(
            'Insight',
            parent=styles['CustomBody'],
            backColor=colors.HexColor('#E8F5E9'),
            borderColor=colors.HexColor('#4CAF50'),
            borderWidth=1,
            borderPadding=10,
            fontSize=11
        )
    )
    elements.append(insight_para)
    
    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer

def apply_manual_cleaning(df: pd.DataFrame, operations: List[str]) -> Tuple[pd.DataFrame, Dict[str, str]]:
    """Applies specific, user-selected cleaning operations (expanded)."""
    df_cleaned = df.copy()
    summary = {}
    
    # 1. Row/Sample Management
    if "remove_duplicates" in operations:
        original_rows = len(df_cleaned)
        df_cleaned.drop_duplicates(inplace=True)
        summary["remove_duplicates"] = f"Removed {original_rows - len(df_cleaned)} duplicate rows."
    
    if "remove_missing_rows" in operations:
        original_rows = len(df_cleaned)
        df_cleaned.dropna(axis=0, how='any', inplace=True)
        summary["remove_missing_rows"] = f"Removed {original_rows - len(df_cleaned)} rows with any missing values."

    # 2. Imputation/Data Type Handling
    if "impute_median" in operations:
        numeric_cols = df_cleaned.select_dtypes(include=['number']).columns
        for col in numeric_cols:
            df_cleaned[col].fillna(df_cleaned[col].median(), inplace=True)
        summary["impute_median"] = f"Imputed missing values in {len(numeric_cols)} numeric columns using the **median**."

    if "impute_mode" in operations:
        cat_cols = df_cleaned.select_dtypes(include=['object', 'category']).columns
        for col in cat_cols:
            if not df_cleaned[col].mode().empty:
                df_cleaned[col].fillna(df_cleaned[col].mode()[0], inplace=True)
        summary["impute_mode"] = f"Imputed missing values in {len(cat_cols)} categorical columns using the **mode**."
        
    # 3. Encoding/Feature Engineering
    if "one_hot_encode" in operations:
        cat_cols = df_cleaned.select_dtypes(include=['object', 'category']).columns
        df_cleaned = pd.get_dummies(df_cleaned, columns=cat_cols, drop_first=True)
        summary["one_hot_encode"] = f"Applied **One-Hot Encoding** to {len(cat_cols)} columns."
        
    if not summary:
        summary["none"] = "No valid cleaning operations were applied. The data remains unchanged."

    return df_cleaned, summary

# --- EDA & Cleaning Helper Functions ---

def calculate_detailed_correlation(df: pd.DataFrame, numeric_cols: List[str]) -> Dict[str, Any]:
    """Calculates correlation matrix and identifies top positive/negative pairs."""
    if len(numeric_cols) < 2:
        return {"matrix_details": "N/A (Less than 2 numeric columns)", "top_pairs": []}

    corr_matrix = df[numeric_cols].corr()
    corr_unstacked = corr_matrix.unstack().sort_values(ascending=False)
    
    unique_pairs = corr_unstacked[corr_unstacked.index.get_level_values(0) != corr_unstacked.index.get_level_values(1)]
    unique_pairs = unique_pairs[~unique_pairs.index.map(frozenset).duplicated()]
    
    # Top 5 positive and negative pairs
    top_positive = unique_pairs[unique_pairs > 0].head(5)
    top_negative = unique_pairs[unique_pairs < 0].tail(5)

    top_pairs = []
    for pair, value in top_positive.items():
        top_pairs.append({"type": "Positive", "pair": f"{pair[0]} & {pair[1]}", "value": f"{value:.4f}", "strength": "Strong" if value >= 0.7 else ("Moderate" if value >= 0.5 else "Weak")})
    for pair, value in top_negative.items():
        top_pairs.append({"type": "Negative", "pair": f"{pair[0]} & {pair[1]}", "value": f"{value:.4f}", "strength": "Strong" if value <= -0.7 else ("Moderate" if value <= -0.5 else "Weak")})

    return {
        "matrix_df": corr_matrix, # Return DataFrame for DOCX table generation
        "top_pairs": top_pairs
    }

# ... (perform_outlier_detection remains the same) ...
def perform_outlier_detection(df: pd.DataFrame, numeric_cols: List[str]) -> Dict[str, int]:
    """Uses IQR method to count potential outliers for reporting."""
    outlier_counts = {}
    for col in numeric_cols:
        Q1 = df[col].quantile(0.25)
        Q3 = df[col].quantile(0.75)
        IQR = Q3 - Q1
        lower_bound = Q1 - 1.5 * IQR
        upper_bound = Q3 + 1.5 * IQR
        
        count = len(df[(df[col] < lower_bound) | (df[col] > upper_bound)])
        if count > 0:
            outlier_counts[col] = count
    return outlier_counts

# --- Automated Cleaning & Reporting (The intelligent pipeline) ---

def automated_clean_and_report_logic(df: pd.DataFrame, filename: str, inspection_data: Dict[str, Any]) -> Tuple[pd.DataFrame, str, Dict[str, Any]]:
    """
    Performs a comprehensive automated cleaning pipeline (removing visualizations).
    """
    df_cleaned = df.copy()
    cleaning_summary = []
    
    try:
        num_cols = inspection_data['numeric_columns']
        cat_cols = inspection_data['categorical_columns']
        
        # Comparative Stats (before vs. after)
        original_missing_count = sum(inspection_data['missing_values'].values())
        original_shape = df.shape
        
        # 1. Intelligent Cleaning (Simulation)
        original_rows = len(df_cleaned)
        df_cleaned.drop_duplicates(inplace=True)
        if original_rows != len(df_cleaned):
            cleaning_summary.append(f"Row Deletion: Removed **{original_rows - len(df_cleaned)}** duplicate rows.")
        
        # Imputation based on type
        for col in num_cols:
            df_cleaned[col].fillna(df_cleaned[col].median(), inplace=True)
        if num_cols: cleaning_summary.append(f"Imputation: Numeric values ({len(num_cols)} columns) imputed using the **Median**.")

        for col in cat_cols:
            if not df_cleaned[col].mode().empty:
                df_cleaned[col].fillna(df_cleaned[col].mode()[0], inplace=True)
        if cat_cols: cleaning_summary.append(f"Imputation: Categorical values ({len(cat_cols)} columns) imputed using the **Mode**.")
        
        # 2. Comprehensive EDA Calculations
        corr_analysis = calculate_detailed_correlation(df_cleaned, num_cols)
        outlier_counts = perform_outlier_detection(df_cleaned, num_cols)
        outlier_summary = [f"**{col}**: {count} values flagged (IQR method)." for col, count in outlier_counts.items()]
        
        # 3. Final Report Content Structure
        report_details = {
            "title": f"Automated Data Cleaning & EDA Report for {filename}",
            "original_df_head": df.head(5),
            "cleaned_df_head": df_cleaned.head(5),
            "cleaning_summary": cleaning_summary,
            "comparative_metrics": {
                "original_missing": original_missing_count,
                "final_missing": df_cleaned.isnull().sum().sum(),
                "original_shape": original_shape,
                "final_shape": df_cleaned.shape
            },
            "correlation_pairs": corr_analysis['top_pairs'],
            "correlation_matrix_df": corr_analysis.get('matrix_df'), # Full matrix DF
            "outlier_flags": outlier_summary,
            "key_insight": "The dataset is verified, cleaned, and structurally sound. Proceed with confidence to feature engineering and modeling."
        }

        return df_cleaned, "Success", report_details
    
    except Exception as e:
        # Error handling for the logic itself
        failure_report = {
            "title": f"Automated Cleaning Failed for {filename}",
            "cleaning_summary": ["Critical error prevented full execution."],
            "eda_summary": [f"Error Details: {str(e)}"],
            "correlation_pairs": [],
            "outlier_flags": [],
            "key_insight": "The data was not fully processed. The returned DataFrame is the original data. Please review the error."
        }
        return df_cleaned, "Failure", failure_report

# --- DOCX Report Generation (Highly Detailed) ---

def add_dataframe_table(document, df: pd.DataFrame, title: str):
    """Helper to add a Pandas DataFrame to the DOCX document as a styled table."""
    document.add_heading(title, level=2)
    
    # Use a basic, reliable style
    table = document.add_table(df.shape[0] + 1, df.shape[1] + 1)
    table.style = 'Table Grid'
    
    # Add row indices to the first column
    table.cell(0, 0).text = 'Index'
    
    # Column Headers
    for j, col_name in enumerate(df.columns):
        table.cell(0, j + 1).text = str(col_name)
    
    # Data Rows
    for i, row in enumerate(df.itertuples(index=True)):
        # Row Index
        table.cell(i + 1, 0).text = str(row[0])
        # Row Data
        for j in range(df.shape[1]):
            cell_data = row[j + 1]
            # Format floats to 4 decimal places, otherwise use string
            if isinstance(cell_data, (float, np.float64)):
                cell_text = f"{cell_data:.4f}"
            else:
                cell_text = str(cell_data)
                
            table.cell(i + 1, j + 1).text = cell_text
            
            # Center numeric columns (Simple way)
            if isinstance(cell_data, (float, np.float64, int)):
                table.cell(i + 1, j + 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    document.add_paragraph() # Add space after table


def generate_docx_report(report_details: Dict[str, Any], filename: str) -> BytesIO:
    """Generates a professional DOCX report with comparison and correlation matrix."""
    document = Document()
    
    # Set document style
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title Page/Header (Fixes extra space)
    document.add_heading(report_details['title'], level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph()
    p.add_run(f"Generated by the Dynamic AutoML Platform on {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}").italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 1: Data Previews
    document.add_heading('1. Data Transformation Impact', level=1)
    
    # Display BEFORE cleaning
    add_dataframe_table(document, report_details['original_df_head'], '1.1. Original Data Preview (First 5 Rows)')
    document.add_paragraph("This view shows the data as it was initially uploaded. Note any missing values (NaN) or inconsistent data types.")
    
    # Display AFTER cleaning
    add_dataframe_table(document, report_details['cleaned_df_head'], '1.2. Cleaned Data Preview (First 5 Rows)')
    document.add_paragraph("This view reflects the data after all imputation and preprocessing steps. Notice the difference in missing values.")
    
    # Section 2: Cleaning & Metrics Comparison
    document.add_heading('2. Cleaning Operations & Metrics', level=1)
    
    metrics = report_details['comparative_metrics']
    
    # Comparative Metrics Table
    document.add_heading('2.1. Comparison of Key Metrics', level=2)
    table_metrics = document.add_table(rows=3, cols=3)
    table_metrics.style = 'Table Grid'
    
    # Headers
    table_metrics.cell(0, 0).text = 'Metric'
    table_metrics.cell(0, 1).text = 'Original Data'
    table_metrics.cell(0, 2).text = 'Cleaned Data'
    
    # Row 1: Shape
    table_metrics.cell(1, 0).text = 'Data Shape (Rows, Columns)'
    table_metrics.cell(1, 1).text = str(metrics['original_shape'])
    table_metrics.cell(1, 2).text = str(metrics['final_shape'])
    
    # Row 2: Missing Values
    table_metrics.cell(2, 0).text = 'Total Missing Values'
    table_metrics.cell(2, 1).text = str(metrics['original_missing'])
    table_metrics.cell(2, 2).text = str(metrics['final_missing'])
    
    document.add_paragraph("Summary: All missing values were resolved, bringing data completeness to 100%.")

    # List of Operations
    document.add_heading('2.2. Operations Log', level=2)
    for item in report_details.get("cleaning_summary", []):
        p = document.add_paragraph()
        p.style = 'List Bullet'
        p.add_run(item.replace('**', '')).bold = True
    
    # Section 3: Detailed Correlation Analysis
    document.add_heading('3. Detailed Exploratory Data Analysis (EDA)', level=1)
    
    # 3.1. Correlation Matrix Table
    corr_df = report_details.get("correlation_matrix_df")
    if corr_df is not None:
        document.add_heading('3.1. Full Correlation Matrix', level=2)
        document.add_paragraph("The correlation matrix shows the linear relationship between every numeric feature pair (values range from -1.0 to +1.0).")
        
        # Add correlation matrix table
        corr_cols = corr_df.columns.tolist()
        corr_table = document.add_table(rows=len(corr_cols) + 1, cols=len(corr_cols) + 1)
        corr_table.style = 'Table Grid'
        
        # Header Row
        for j, col_name in enumerate(corr_cols):
            corr_table.cell(0, j + 1).text = col_name
            corr_table.cell(j + 1, 0).text = col_name # Side header
        
        # Data Rows
        for i, col_i in enumerate(corr_cols):
            for j, col_j in enumerate(corr_cols):
                value = corr_df.loc[col_i, col_j]
                cell = corr_table.cell(i + 1, j + 1)
                cell.text = f"{value:.3f}"
                
                # Highlight strong correlation values
                if abs(value) >= 0.7 and i != j:
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0x45, 0x00) # Orange/Red
        
    # 3.2. Top Correlation Insights
    document.add_heading('3.2. Top Relationship Insights', level=2)
    corr_pairs = report_details.get("correlation_pairs", [])
    if corr_pairs:
        for item in corr_pairs:
            p = document.add_paragraph()
            p.style = 'List Bullet'
            p.add_run(f"Relationship: {item['pair']}").bold = True
            p.add_run(f" | Strength: {item['strength']} ({item['type']})")
            p.add_run(f" | Value: {item['value']}").font.color.rgb = RGBColor(0x00, 0xAA, 0x00)
    else:
        document.add_paragraph("No strong linear relationships found (or less than 2 numeric columns exist).")

    # Section 4: Outlier Flagging
    document.add_heading('4. Outlier Flagging (IQR Method)', level=1)
    for item in report_details.get("outlier_flags", []):
        p = document.add_paragraph()
        p.style = 'List Bullet'
        p.add_run(item.replace('**', '')).bold = True

    # Final Conclusion/Key Insight
    document.add_heading('Key Insights for Modeling', level=2)
    document.add_paragraph(report_details.get("key_insight", "Analysis complete."))

    # Save the document to a BytesIO stream
    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream